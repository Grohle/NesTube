"""
nestube/nestube/nestube/export_utils.py
PDF and Excel export, colour-mapping for nesting visualisation.
Uses Unicode TTF fonts for PDF generation.
"""
from __future__ import annotations

import os
import random
import json
import subprocess
import tempfile
import sys
try:  # tkinter is only needed by the legacy Tk-dialog export helpers
    from tkinter import filedialog, messagebox
except ImportError:  # Qt-only / headless: the pure export functions still work
    filedialog = messagebox = None
from typing import Dict, List, Optional, Tuple

import pandas as pd
from fpdf import FPDF

from .models import AppState
from .i18n import t


def _open_file(path: str) -> None:
    """Open a file with the default OS application (cross-platform)."""
    try:
        if sys.platform == "win32":
            os.startfile(path)
        elif sys.platform == "darwin":
            subprocess.Popen(["open", path])
        else:
            subprocess.Popen(["xdg-open", path])
    except Exception:
        pass
from . import app_config
from . import units
from .logic import calcular_resultado


# ── Colour mapping ────────────────────────────────────────────────────────────

RETAL_COLOR: str = "#39FF14"  # neon green — RESERVED for remnants only

_PALETTE: List[str] = [
    "#4E9AF1", "#E8635A", "#F2C94C", "#BB87F7", "#56CCF2",
    "#F2994A", "#EB5757", "#9B51E0", "#2F80ED", "#FF6F91",
    "#1F9EFF", "#D65DB1", "#FF9671", "#FFC75F", "#C08070",
    "#4CC9F0", "#7B2FBE", "#E07A5F", "#3D405B", "#C1A057",
    "#5E81F4", "#F47C87", "#54C6EB", "#F6AE2D", "#A23B72",
    "#B07DAB", "#E84855", "#8B5CF6", "#F59E0B", "#D97706",
    "#EF4444", "#3B82F6", "#EC4899", "#7C6BAF", "#F97316",
    "#6366F1",
]
_color_cache: dict = {}
_palette_index: int = 0

_PDF_FONT = "NesTubeUnicode"


def _is_too_close_to_retal(hex_color: str) -> bool:
    """True if a hex color is perceptually close to neon green (#39FF14)."""
    r = int(hex_color[1:3], 16)
    g = int(hex_color[3:5], 16)
    b = int(hex_color[5:7], 16)
    return g > 160 and g > r + 40 and g > b + 40


def get_cut_color(key) -> str:
    """Return a stable color for a cut key; never returns RETAL_COLOR."""
    global _palette_index
    if key not in _color_cache:
        if _palette_index < len(_PALETTE):
            _color_cache[key] = _PALETTE[_palette_index]
            _palette_index += 1
        else:
            h = hash(str(key)) & 0xFFFFFF
            candidate = "#{:06x}".format(h)
            for _ in range(10):
                if not _is_too_close_to_retal(candidate):
                    break
                h = (h + 0x801020) & 0xFFFFFF
                candidate = "#{:06x}".format(h)
            _color_cache[key] = candidate
    return _color_cache[key]


def get_retal_color() -> str:
    """Return the reserved neon-green color for remnant bar segments."""
    return RETAL_COLOR


def clear_color_cache() -> None:
    global _palette_index
    _color_cache.clear()
    _palette_index = 0


# ── PDF font helpers ──────────────────────────────────────────────────────────

def _resolve_pdf_font_paths() -> Tuple[str, str]:
    prefs = app_config.get()
    regular = prefs.pdf_font_regular
    bold = prefs.pdf_font_bold or regular
    if not regular or not os.path.isfile(regular):
        for path in app_config._discover_pdf_font_paths():  # noqa: SLF001
            regular = path
            bold_candidate = path.replace("Sans.ttf", "Sans-Bold.ttf")
            if not os.path.isfile(bold_candidate):
                bold_candidate = path.replace(".ttf", "bd.ttf")
            if not os.path.isfile(bold_candidate):
                bold_candidate = path
            bold = bold_candidate
            prefs.pdf_font_regular = regular
            prefs.pdf_font_bold = bold
            app_config.save(prefs)
            break
    if not regular:
        raise FileNotFoundError(
            "No Unicode TTF font found for PDF export. "
            "Configure a font in Settings → PDF font."
        )
    if not os.path.isfile(bold):
        bold = regular
    return regular, bold


def _register_pdf_fonts(pdf: FPDF) -> None:
    regular, bold = _resolve_pdf_font_paths()
    pdf.add_font(_PDF_FONT, "", regular)
    if bold != regular:
        pdf.add_font(_PDF_FONT, "B", bold)
    else:
        pdf.add_font(_PDF_FONT, "B", regular)


def set_pdf_font(pdf: FPDF, style: str = "", size: int = 10) -> None:
    pdf.set_font(_PDF_FONT, style, size)


# ── Helpers ───────────────────────────────────────────────────────────────────

def _next_number(directory: str, ext: str) -> int:
    try:
        files = [f for f in os.listdir(directory) if f.endswith(f".{ext}")]
    except OSError:
        return 1
    nums = [
        int(f.split("_")[-1].split(".")[0])
        for f in files
        if f.split("_")[-1].split(".")[0].isdigit()
    ]
    return max(nums, default=0) + 1


# ── Excel export ──────────────────────────────────────────────────────────────

def exportar_excel(state: AppState, filename: Optional[str] = None) -> None:
    if not state.cortes:
        if filename is None and messagebox:
            messagebox.showerror(t("export_error"), t("no_data_msg"))
        else:
            raise ValueError(t("no_data_msg"))
        return

    if filename is None:
        initial_dir = state.export_path or "."
        filename = filedialog.asksaveasfilename(
            defaultextension=".xlsx",
            filetypes=[("Excel", "*.xlsx")],
            initialfile=f"NesTube_Resultados_{_next_number(initial_dir, 'xlsx')}.xlsx",
            initialdir=initial_dir,
        )
        if not filename:
            return

    n_ingletes = sum(1 for c in state.cortes if c.inglete1 or c.inglete2)
    rows = []
    total_kg = 0.0
    total_cost = 0.0
    for c in state.cortes:
        res = calcular_resultado(c, state.perfil, state.barras_necesarias, state.longitud_barra, n_ingletes)
        line = {
            t("description"): c.descripcion,
            t("placeholder_length", u=units.u_len()): c.largo,
            t("placeholder_qty"): c.cantidad,
            "Kg/ud": round(res.kg_ud, 4),
            "m²/ud": round(res.m2_ud, 6),
            f"{t('material_section')}/ud ({state.currency})": round(res.precio_material_ud, 2),
            f"{t('labour_section')}/ud ({state.currency})": round(res.coste_mano_obra_ud, 2),
            f"Total/ud ({state.currency})": round(res.precio_total_ud, 2),
            f"Total ({state.currency})": round(res.precio_total_linea, 2),
        }
        total_kg += res.kg_ud * c.cantidad
        total_cost += res.precio_total_linea
        rows.append(line)
    rows.append({
        t("description"): "TOTAL",
        f"Total ({state.currency})": round(total_cost, 2),
    })

    with pd.ExcelWriter(filename, engine="openpyxl") as writer:
        header_data = {
            t("description"): [state.descripcion or ""],
            t("order_number"): [state.pedido or ""],
            t("offer_number"): [state.oferta or ""],
            t("client"): [state.cliente or ""],
        }
        for name, value in state.custom_fields.items():
            if value:
                header_data[name] = [value]
        header_data["Total Kg"] = [round(total_kg, 3)]
        header_data[f"Total ({state.currency})"] = [round(total_cost, 2)]
        if state.barras_necesarias:
            from .logic import eficiencia_barras
            ef = eficiencia_barras(state.barras_necesarias, state.longitud_barra)
            header_data[t("bars")] = [len(state.barras_necesarias)]
            header_data[t("efficiency")] = [f"{ef:.1f}%"]

        pd.DataFrame(header_data).to_excel(writer, sheet_name=t("description"), index=False)
        pd.DataFrame(rows).to_excel(writer, sheet_name=t("cut_list"), index=False)

    _open_file(filename)


# ── PDF export ────────────────────────────────────────────────────────────────

def should_offer_multi_material_pdf(state: AppState) -> bool:
    """True when other material contexts have nesting data worth offering in PDF."""
    from .context_sync import ensure_material_contexts, effective_barras

    ensure_material_contexts(state)
    if len(state.material_contexts) <= 1:
        return False
    active = state.active_material_index
    return any(
        effective_barras(ctx)
        for i, ctx in enumerate(state.material_contexts)
        if i != active
    )


def ask_include_other_materials() -> Optional[bool]:
    """Ask whether to include other materials in the PDF export.

    Returns True (all materials), False (current only), or None (cancelled).
    """
    import tkinter as tk

    result: List[Optional[bool]] = [None]

    dlg = tk.Toplevel()
    dlg.title(t("export_pdf_btn"))
    dlg.transient()
    dlg.grab_set()
    dlg.resizable(False, False)

    tk.Label(
        dlg,
        text=t("pdf_ask_other_materials"),
        wraplength=360,
        justify="center",
        padx=16,
        pady=12,
    ).pack()

    btn_row = tk.Frame(dlg)
    btn_row.pack(pady=(0, 14))

    def _choose(include_all: bool) -> None:
        result[0] = include_all
        dlg.destroy()

    tk.Button(
        btn_row,
        text=t("pdf_current_only"),
        width=22,
        command=lambda: _choose(False),
    ).pack(side=tk.LEFT, padx=6)
    tk.Button(
        btn_row,
        text=t("pdf_all_materials"),
        width=22,
        command=lambda: _choose(True),
    ).pack(side=tk.LEFT, padx=6)

    dlg.protocol("WM_DELETE_WINDOW", dlg.destroy)
    dlg.wait_window()
    return result[0]


def exportar_pdf(
    state: AppState,
    prefix: str = "NesTube_Anidado",
    bars_override: List[List[float]] | None = None,
    material_contexts: list | None = None,
    filename: Optional[str] = None,
) -> None:
    """Export nesting PDF.

    bars_override: if provided, export only these bars instead of state.barras_necesarias.
    material_contexts: if provided, one PDF section per material context.
    filename: if provided, skip the file-chooser dialog (Qt callers supply this).
    """
    if material_contexts:
        from .context_sync import effective_barras

        has_data = any(effective_barras(ctx) for ctx in material_contexts)
        if not has_data:
            raise ValueError(t("no_nesting_data"))
    else:
        bars = bars_override if bars_override is not None else state.barras_necesarias
        if not bars:
            raise ValueError(t("no_nesting_data"))

    if filename is None:
        initial_dir = state.export_path or "."
        filename = filedialog.asksaveasfilename(
            defaultextension=".pdf",
            filetypes=[("PDF", "*.pdf")],
            initialfile=f"{prefix}_{_next_number(initial_dir, 'pdf')}.pdf",
            initialdir=initial_dir,
        )
        if not filename:
            return

    if material_contexts:
        _write_pdf_all_materials(filename, state, material_contexts)
    else:
        _write_pdf(filename, state, bars_override=bars_override)
    _open_file(filename)


def _pdf_hatch(pdf: FPDF, x: float, y: float, w: float, h: float,
               spacing: float = 2.0) -> None:
    """Draw diagonal hatching lines in a PDF rectangle."""
    pdf.set_draw_color(180, 180, 185)
    total = w + h
    offset = spacing
    while offset < total:
        lx0 = x + min(offset, w)
        ly0 = y + max(0, offset - w)
        lx1 = x + max(0, offset - h)
        ly1 = y + min(offset, h)
        pdf.line(lx0, ly0, lx1, ly1)
        offset += spacing


def _write_pdf(
    filename: str,
    state: AppState,
    bars_override: List[List[float]] | None = None,
) -> None:
    bars = bars_override if bars_override is not None else state.barras_necesarias
    pdf = FPDF(orientation="P", unit="mm", format="A4")
    _register_pdf_fonts(pdf)
    pdf.add_page()

    margin = 12
    page_w = 210 - 2 * margin
    bar_h = 9
    bar_gap = 14
    y = 12

    _write_pdf_header(pdf, state)

    pdf.ln(4)
    y = pdf.get_y() + 2

    set_pdf_font(pdf, "", 7)
    kerf = max(state.perdida_corte, 0.0)
    if state.longitud_barra <= 0:
        raise ValueError(t("invalid_bar_length"))
    for i, barra in enumerate(bars):
        retal = state.longitud_barra - sum(barra)
        pdf.set_xy(margin, y)
        set_pdf_font(pdf, "B", 8)
        pdf.cell(
            0, 7,
            t("pdf_bar_info",
              n=str(i + 1),
              pieces=str(len(barra)),
              retal=f"{retal:.1f}",
              usage=f"{sum(barra)/state.longitud_barra*100:.1f}"),
            0, 1,
        )
        y += 7
        x = margin
        set_pdf_font(pdf, "", 6)
        for idx, corte in enumerate(barra):
            kerf_w = (kerf / state.longitud_barra) * page_w if idx < len(barra) - 1 else 0
            cw = max((corte / state.longitud_barra) * page_w - kerf_w, 0.5)
            color = get_cut_color(corte)
            r, g, b = int(color[1:3], 16), int(color[3:5], 16), int(color[5:], 16)
            pdf.set_fill_color(r, g, b)
            pdf.rect(x, y, cw, bar_h, "DF")
            if cw > 14:
                pdf.set_xy(x + 0.5, y + 1.5)
                txt = f"{corte:.0f}"
                pdf.cell(cw - 1, bar_h - 2, txt, 0, 0, "C")
            x += cw
            if kerf_w > 0:
                pdf.set_draw_color(80, 80, 84)
                pdf.line(x, y + 1, x, y + bar_h - 1)
                x += kerf_w

        if retal > 0.5:
            rw = (retal / state.longitud_barra) * page_w
            pdf.set_draw_color(160, 160, 165)
            pdf.rect(x, y, rw, bar_h, "D")
            _pdf_hatch(pdf, x, y, rw, bar_h)
        y += bar_h + bar_gap

        if y > 270:
            pdf.add_page()
            y = 15

    pdf.add_page()
    set_pdf_font(pdf, "B", 12)
    pdf.cell(0, 10, t("pdf_legend_title"), 0, 1, "C")
    set_pdf_font(pdf, "", 9)
    pdf.ln(4)

    legend: Dict[float, int] = {}
    for barra in bars:
        for c in barra:
            legend[c] = legend.get(c, 0) + 1

    for corte, qty in sorted(legend.items(), reverse=True):
        color = get_cut_color(corte)
        r, g, b = int(color[1:3], 16), int(color[3:5], 16), int(color[5:], 16)
        pdf.set_fill_color(r, g, b)
        cy = pdf.get_y() + 1
        pdf.rect(margin, cy, 10, 7, "DF")
        pdf.set_xy(margin + 13, pdf.get_y())
        pdf.cell(0, 9, f"{corte:.2f} mm   x   {qty} {t('pieces')}", 0, 1)
        pdf.ln(1)

    pdf.output(filename)


def _write_pdf_all_materials(filename: str, state: AppState, contexts: list) -> None:
    """One PDF file with a section per material context."""
    from .context_sync import effective_barras

    pdf = FPDF(orientation="P", unit="mm", format="A4")
    _register_pdf_fonts(pdf)
    margin = 12
    wrote_any = False
    for ctx in contexts:
        bars = effective_barras(ctx)
        if not bars:
            continue
        wrote_any = True
        pdf.add_page()
        section = AppState.from_dict({
            **state.to_dict(),
            "descripcion": ctx.display_name or ctx.material or ctx.name,
            "calidad": ctx.quality,
            "longitud_barra": ctx.longitud_barra,
            "perdida_corte": ctx.perdida_corte,
            "margen_tubo": ctx.margen_tubo,
            "barras_necesarias": bars,
        })
        _write_pdf_header(pdf, section)
        pdf.ln(4)
        y = pdf.get_y() + 2
        page_w = 210 - 2 * margin
        bar_h = 9
        bar_gap = 14
        kerf = max(section.perdida_corte, 0.0)
        if section.longitud_barra <= 0:
            continue
        set_pdf_font(pdf, "", 7)
        for i, barra in enumerate(bars):
            retal = section.longitud_barra - sum(barra)
            pdf.set_xy(margin, y)
            set_pdf_font(pdf, "B", 8)
            pdf.cell(
                0, 7,
                t("pdf_bar_info",
                  n=str(i + 1),
                  pieces=str(len(barra)),
                  retal=f"{retal:.1f}",
                  usage=f"{sum(barra)/section.longitud_barra*100:.1f}"),
                0, 1,
            )
            y += 7
            x = margin
            set_pdf_font(pdf, "", 6)
            for idx, corte in enumerate(barra):
                kerf_w = (kerf / section.longitud_barra) * page_w if idx < len(barra) - 1 else 0
                cw = max((corte / section.longitud_barra) * page_w - kerf_w, 0.5)
                color = get_cut_color(corte)
                r, g, b = int(color[1:3], 16), int(color[3:5], 16), int(color[5:], 16)
                pdf.set_fill_color(r, g, b)
                pdf.rect(x, y, cw, bar_h, "DF")
                x += cw
                if kerf_w > 0:
                    pdf.set_draw_color(80, 80, 84)
                    pdf.line(x, y + 1, x, y + bar_h - 1)
                    x += kerf_w
            if retal > 0.5:
                rw = (retal / section.longitud_barra) * page_w
                pdf.set_draw_color(160, 160, 165)
                pdf.rect(x, y, rw, bar_h, "D")
                _pdf_hatch(pdf, x, y, rw, bar_h)
            y += bar_h + bar_gap
            if y > 270:
                pdf.add_page()
                y = 15
    if not wrote_any:
        raise ValueError(t("no_nesting_data"))
    pdf.output(filename)


def _write_pdf_header(pdf: FPDF, state: AppState) -> None:
    set_pdf_font(pdf, "B", 14)
    pdf.cell(0, 9, t("pdf_header"), 0, 1, "C")
    set_pdf_font(pdf, "", 9)
    layout = _load_layout()
    values = {
        "description": state.descripcion,
        "order_number": state.pedido,
        "offer_number": state.oferta,
        "client": state.cliente,
    }
    if layout:
        x0 = 12
        y0 = pdf.get_y() + 2
        for field, pos in layout.items():
            val = values.get(field, "")
            if not val and field != "custom_fields":
                continue
            x = x0 + float(pos.get("x", 0)) * 0.2
            y = y0 + float(pos.get("y", 0)) * 0.2
            pdf.set_xy(x, y)
            if field == "custom_fields":
                txt = " · ".join(f"{k}: {v}" for k, v in state.custom_fields.items() if v)
            elif field == "order_number":
                txt = f"{t('header_order_number')}: {values['order_number']}"
            elif field == "offer_number":
                txt = f"{t('header_offer_number')}: {values['offer_number']}"
            else:
                txt = f"{field}: {val}"
            pdf.cell(0, 5, txt, 0, 1)
    else:
        if state.descripcion:
            pdf.cell(0, 6, t("pdf_description", desc=state.descripcion), 0, 1, "C")
        pdf.cell(
            0, 6,
            t("pdf_order_client", order=state.pedido, client=state.cliente),
            0, 1, "C",
        )
        if state.oferta:
            pdf.cell(0, 6, f"{t('header_offer_number')}: {state.oferta}", 0, 1, "C")
        if state.custom_fields:
            for name, value in state.custom_fields.items():
                if value:
                    pdf.cell(0, 5, f"{name}: {value}", 0, 1, "C")
    pdf.ln(4)


def _load_layout() -> dict:
    path = app_config.get().pdf_template_layout_path
    if not path or not os.path.isfile(path):
        return {}
    try:
        with open(path, "r", encoding="utf-8") as fh:
            return json.load(fh)
    except (OSError, ValueError):
        return {}


def exportar_presupuesto_pdf(state: AppState, filename: Optional[str] = None) -> None:
    """Export a formal budget/quote PDF with costs, client data, and cut details."""
    if not state.cortes:
        raise ValueError(t("no_cuts_msg"))

    if filename is None:
        initial_dir = state.export_path or "."
        filename = filedialog.asksaveasfilename(
            defaultextension=".pdf",
            filetypes=[("PDF", "*.pdf")],
            initialfile=f"NesTube_Presupuesto_{_next_number(initial_dir, 'pdf')}.pdf",
            initialdir=initial_dir,
        )
        if not filename:
            return

    _write_presupuesto_pdf(filename, state)
    _open_file(filename)


def _write_presupuesto_pdf(filename: str, state: AppState) -> None:
    """Generate a formal quote/order PDF."""
    pdf = FPDF(orientation="P", unit="mm", format="A4")
    _register_pdf_fonts(pdf)
    pdf.add_page()

    set_pdf_font(pdf, "B", 16)
    pdf.cell(0, 10, t("pdf_header"), 0, 1, "C")
    pdf.ln(2)

    set_pdf_font(pdf, "", 10)
    if state.descripcion:
        pdf.cell(0, 6, f"{t('description')}: {state.descripcion}", 0, 1)
    if state.pedido:
        pdf.cell(0, 6, f"{t('order_number')}: {state.pedido}", 0, 1)
    if state.oferta:
        pdf.cell(0, 6, f"{t('offer_number')}: {state.oferta}", 0, 1)
    if state.cliente:
        pdf.cell(0, 6, f"{t('client')}: {state.cliente}", 0, 1)
    for name, value in state.custom_fields.items():
        if value:
            pdf.cell(0, 6, f"{name}: {value}", 0, 1)
    pdf.ln(4)

    set_pdf_font(pdf, "B", 9)
    col_w = [10, 50, 22, 16, 22, 22, 22, 22]
    headers = ["#", t("description"), t("placeholder_length", u=units.u_len()),
               t("placeholder_qty"), "Kg/ud", f"{t('material_section')}",
               f"{t('labour_section')}", "Total"]
    for i, h in enumerate(headers):
        pdf.cell(col_w[i], 7, h, 1, 0, "C")
    pdf.ln()

    set_pdf_font(pdf, "", 8)
    n_ingletes = sum(1 for c in state.cortes if c.inglete1 or c.inglete2)
    total_cost = 0.0
    total_kg = 0.0
    for idx, c in enumerate(state.cortes):
        res = calcular_resultado(c, state.perfil, state.barras_necesarias,
                                 state.longitud_barra, n_ingletes)
        total_cost += res.precio_total_linea
        total_kg += res.kg_ud * c.cantidad
        row = [
            str(idx + 1),
            c.descripcion[:25] or f"Cut {idx+1}",
            f"{c.largo:.0f}",
            str(c.cantidad),
            f"{res.kg_ud:.3f}",
            f"{res.precio_material_ud:.2f}",
            f"{res.coste_mano_obra_ud:.2f}",
            f"{res.precio_total_linea:.2f}",
        ]
        for i, cell in enumerate(row):
            pdf.cell(col_w[i], 6, cell, 1, 0, "C" if i > 1 else "L")
        pdf.ln()

        if pdf.get_y() > 265:
            pdf.add_page()
            set_pdf_font(pdf, "", 8)

    pdf.ln(2)
    set_pdf_font(pdf, "B", 10)
    pdf.cell(0, 8, f"{t('total_order', total=f'{total_cost:,.2f}', currency=state.currency)}", 0, 1)
    set_pdf_font(pdf, "", 9)
    pdf.cell(0, 6, f"{t('weight_per_unit')}: {total_kg:.2f} kg", 0, 1)

    if state.barras_necesarias:
        pdf.ln(4)
        set_pdf_font(pdf, "B", 9)
        from .logic import eficiencia_barras
        ef = eficiencia_barras(state.barras_necesarias, state.longitud_barra)
        pdf.cell(0, 6, f"{t('nesting_title')}: {len(state.barras_necesarias)} {t('bars')} "
                 f"- {t('efficiency')} {ef:.1f}%", 0, 1)

    pdf.output(filename)


def exportar_docx(state: AppState, filename: Optional[str] = None) -> None:
    """Export a formal quote/order as editable DOCX."""
    if not state.cortes:
        raise ValueError(t("no_cuts_msg"))

    if filename is None:
        initial_dir = state.export_path or "."
        filename = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word DOCX", "*.docx")],
            initialfile=f"NesTube_Presupuesto_{_next_number(initial_dir, 'docx')}.docx",
            initialdir=initial_dir,
        )
        if not filename:
            return

    _write_docx(filename, state)
    _open_file(filename)


def _write_docx(filename: str, state: AppState) -> None:
    """Generate a formal quote/order DOCX document."""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()
    style = doc.styles['Normal']
    style.font.size = Pt(10)

    title = doc.add_heading(t("pdf_header"), level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if state.descripcion:
        doc.add_paragraph(f"{t('description')}: {state.descripcion}")
    if state.pedido:
        doc.add_paragraph(f"{t('order_number')}: {state.pedido}")
    if state.oferta:
        doc.add_paragraph(f"{t('offer_number')}: {state.oferta}")
    if state.cliente:
        doc.add_paragraph(f"{t('client')}: {state.cliente}")
    for name, value in state.custom_fields.items():
        if value:
            doc.add_paragraph(f"{name}: {value}")

    doc.add_paragraph()

    headers = ["#", t("description"), t("placeholder_length", u=units.u_len()),
               t("placeholder_qty"), "Kg/ud",
               f"{t('material_section')} ({state.currency})",
               f"{t('labour_section')} ({state.currency})",
               f"Total ({state.currency})"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(8)

    n_ingletes = sum(1 for c in state.cortes if c.inglete1 or c.inglete2)
    total_cost = 0.0
    total_kg = 0.0
    for idx, c in enumerate(state.cortes):
        res = calcular_resultado(c, state.perfil, state.barras_necesarias,
                                 state.longitud_barra, n_ingletes)
        total_cost += res.precio_total_linea
        total_kg += res.kg_ud * c.cantidad
        row_cells = table.add_row().cells
        values = [
            str(idx + 1),
            c.descripcion or f"Cut {idx+1}",
            f"{c.largo:.0f}",
            str(c.cantidad),
            f"{res.kg_ud:.3f}",
            f"{res.precio_material_ud:.2f}",
            f"{res.coste_mano_obra_ud:.2f}",
            f"{res.precio_total_linea:.2f}",
        ]
        for i, val in enumerate(values):
            row_cells[i].text = val
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    doc.add_paragraph()
    total_para = doc.add_paragraph()
    total_run = total_para.add_run(
        t("total_order", total=f"{total_cost:,.2f}", currency=state.currency)
    )
    total_run.bold = True
    total_run.font.size = Pt(12)

    doc.add_paragraph(f"{t('weight_per_unit')}: {total_kg:.2f} kg")

    if state.barras_necesarias:
        from .logic import eficiencia_barras
        ef = eficiencia_barras(state.barras_necesarias, state.longitud_barra)
        doc.add_paragraph(
            f"{t('nesting_title')}: {len(state.barras_necesarias)} {t('bars')} "
            f"- {t('efficiency')} {ef:.1f}%"
        )

    doc.save(filename)


def imprimir(
    state: AppState,
    bars_override: List[List[float]] | None = None,
) -> None:
    """Generate a temporary PDF and send it to the system printer."""
    bars = bars_override if bars_override is not None else state.barras_necesarias
    if not bars and not state.cortes:
        raise ValueError(t("no_data_msg"))

    tmp_path: Optional[str] = None
    try:
        fd, tmp_path = tempfile.mkstemp(suffix=".pdf")
        os.close(fd)

        if bars:
            _write_pdf(tmp_path, state, bars_override=bars)
        else:
            _write_presupuesto_pdf(tmp_path, state)

        if sys.platform == "win32":
            # Async hand-off to the associated print app — it reads the file
            # after this returns, so deleting now would truncate the job. Delete
            # later on a daemon timer instead.
            os.startfile(tmp_path, "print")
            import threading

            def _later(p=tmp_path):
                try:
                    os.unlink(p)
                except OSError:
                    pass
            threading.Timer(60.0, _later).start()
        else:
            # lp/lpr COPY the file into the spool queue before exiting, so block
            # until they finish (run, not Popen) and only then delete the temp.
            cmd = "lpr" if sys.platform == "darwin" else "lp"
            try:
                subprocess.run([cmd, tmp_path], check=False)
            finally:
                try:
                    os.unlink(tmp_path)
                except OSError:
                    pass
    except Exception:
        # On failure before hand-off, don't leak the temp file.
        if tmp_path and os.path.isfile(tmp_path):
            try:
                os.unlink(tmp_path)
            except OSError:
                pass
        raise
